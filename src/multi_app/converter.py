from PyPDF2 import PdfReader, PdfWriter
from pdf2image import convert_from_path
from PIL import Image
import pytesseract
from docx import Document
import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph


def extract_text_from_pdf(pdf_path):
    with open(pdf_path, 'rb') as file:
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text

def convert_pdf_to_images(pdf_path, output_folder):
    images = convert_from_path(pdf_path)
    for i, image in enumerate(images):
        image.save(f"{output_folder}/page_{i}.png", 'PNG')

def convert_image_to_text(image_path):
    image = Image.open(image_path)
    text = pytesseract.image_to_string(image)
    return text

def convert_word_to_text(word_path):
    doc = Document(word_path)
    text = "\n".join([para.text for para in doc.paragraphs])  # Extrai o texto dos parágrafos
    return text

def convert_excel_to_text(excel_path):
    df = pd.read_excel(excel_path)  # Lê o arquivo Excel
    return df.to_string()  # Converte o conteúdo do DataFrame para uma string legível

def remove_pdf_pages(pdf_path, pages_to_remove, output_path):
    reader = PdfReader(pdf_path)
    writer = PdfWriter()
    
    for i in range(len(reader.pages)):
        if i not in pages_to_remove:
            writer.add_page(reader.pages[i])
    
    with open(output_path, 'wb') as output_file:
        writer.write(output_file)

def convert_word_to_pdf(word_path, pdf_path):
    doc = Document(word_path)
    # Criação do documento PDF com ReportLab
    pdf = SimpleDocTemplate(pdf_path, pagesize=letter)
    elements = []

    # Estilo para o texto (Fonte padrão)
    style = getSampleStyleSheet()['Normal']
    
    # Extrair parágrafos do Word e converter para Parágrafos do ReportLab
    for para in doc.paragraphs:
        paragraph_text = para.text.strip()
        if paragraph_text:  # Não adicionar parágrafos vazios
            paragraph = Paragraph(paragraph_text, style)
            elements.append(paragraph)

    # Gerar o PDF com o conteúdo extraído do Word
    pdf.build(elements)
    print(f"PDF salvo em: {pdf_path}")